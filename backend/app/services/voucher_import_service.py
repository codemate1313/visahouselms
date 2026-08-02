from __future__ import annotations

import io
import re
import zipfile
from typing import List, Set
from fastapi import HTTPException, UploadFile, status

# Regex for matching exact 16-digit alphanumeric codes (e.g. ABCD1234EFGH5678 or formatted ABCD-1234-EFGH-5678)
RAW_CODE_REGEX = re.compile(r'\b[A-Za-z0-9]{16}\b')
FORMATTED_CODE_REGEX = re.compile(r'\b[A-Za-z0-9]{4}-[A-Za-z0-9]{4}-[A-Za-z0-9]{4}-[A-Za-z0-9]{4}\b')


def extract_voucher_codes_from_bytes(filename: str, content: bytes) -> List[str]:
    """Extract unique 16-digit alphanumeric voucher codes from file bytes.
    Supports PDF (.pdf), Excel (.xlsx, .xls, .csv), Word (.docx), and Text (.txt) files.
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    extracted_text_blocks: List[str] = []

    if ext == "pdf":
        extracted_text_blocks = _extract_from_pdf(content)
    elif ext in ("xlsx", "xls", "csv"):
        extracted_text_blocks = _extract_from_excel_or_csv(ext, content)
    elif ext == "docx":
        extracted_text_blocks = _extract_from_docx(content)
    elif ext in ("txt", "log", "text"):
        extracted_text_blocks = [content.decode("utf-8", errors="ignore")]
    else:
        # Fallback to UTF-8 decoding
        extracted_text_blocks = [content.decode("utf-8", errors="ignore")]

    found_codes: Set[str] = set()

    for block in extracted_text_blocks:
        if not block:
            continue
        # 1. Match formatted codes (XXXX-XXXX-XXXX-XXXX)
        for fmt in FORMATTED_CODE_REGEX.findall(block):
            clean = fmt.replace("-", "").strip().upper()
            if len(clean) == 16:
                found_codes.add(clean)

        # 2. Match raw 16-digit alphanumeric codes
        for raw in RAW_CODE_REGEX.findall(block):
            clean = raw.strip().upper()
            if len(clean) == 16:
                found_codes.add(clean)

    return list(found_codes)


def _extract_from_pdf(content: bytes) -> List[str]:
    blocks: List[str] = []
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        for page in reader.pages:
            text = page.extract_text()
            if text:
                blocks.append(text)
    except Exception as exc:
        # Fallback raw string search if PDF parsing encounters an error
        blocks.append(content.decode("latin-1", errors="ignore"))
    return blocks


def _extract_from_excel_or_csv(ext: str, content: bytes) -> List[str]:
    blocks: List[str] = []
    if ext == "csv":
        blocks.append(content.decode("utf-8", errors="ignore"))
    else:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                for row in sheet.iter_rows(values_only=True):
                    row_str = " ".join([str(cell) for cell in row if cell is not None])
                    if row_str:
                        blocks.append(row_str)
        except Exception:
            blocks.append(content.decode("latin-1", errors="ignore"))
    return blocks


def _extract_from_docx(content: bytes) -> List[str]:
    blocks: List[str] = []
    # Method A: python-docx
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        for para in doc.paragraphs:
            if para.text:
                blocks.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        blocks.append(cell.text)
        if blocks:
            return blocks
    except Exception:
        pass

    # Method B: Direct zipfile extract of word/document.xml
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            if "word/document.xml" in z.namelist():
                xml_content = z.read("word/document.xml").decode("utf-8", errors="ignore")
                # Strip XML tags
                clean_xml = re.sub(r'<[^>]+>', ' ', xml_content)
                blocks.append(clean_xml)
    except Exception:
        blocks.append(content.decode("latin-1", errors="ignore"))

    return blocks
